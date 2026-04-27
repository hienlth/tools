"""
Sinh BẢNG ĐỀ NGHỊ THANH TOÁN từ Data.xlsx và Template.docx
Yêu cầu: pip install python-docx openpyxl lxml

Cấu trúc cột Data.xlsx (bắt đầu từ dòng 3):
  col 0: Mã GV
  col 1: Họ tên GV
  col 2: Trừ thuế TNCN  ("X" = trừ 10%)
  col 3: Tự luận        (số bài)
  col 4: Thực hành      (số bài)
  col 5: Tiểu luận      (số bài)
  col 6: Ra đề tự luận  (số đề)
  col 7: Duyệt đề tự luận (số đề)

Cấu trúc header Data.xlsx:
  D1: Năm học  (vd: "2025 - 2026")
  F1: Học kỳ   (vd: "1")

"""

import copy, io
import openpyxl
import lxml.etree as etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =================== CẤU HÌNH GIÁ ===================
GIA = {
    "tu_luan":   3000,
    "thuc_hanh": 1800,
    "tieu_luan": 3000,
    "ra_de":     90000,
    "duyet_de":  30000,
    "huong_dan_khoa_luan": 800000, # không quá 5
    "chu_tich_khoa_luan": 150000,
    "phan_bien_khoa_luan": 200000,
    "thu_ky_khoa_luan": 120000,
    "huong_dan_ho_so_tn": 1000000,
    "cham_ho_so_tn": 25000,
    "cham_tieu_luan_nckh": 100000,
    "huong_dan_tieu_luan_nckh": 400000, # không quá 10
    "cham_nckh_tich_luy": 100000,
}

ACT_NAME = {
    "ra_de":     "Ra đề và đáp án thi viết 90-120 phút",
    "duyet_de":  "Duyệt đề và ĐA thi viết 90-120 phút",
    "tu_luan":   "Chấm bài thi hết HP tự luận bậc ĐH KHTN",
    "thuc_hanh": "Chấm bài thi thực hành bậc ĐH KHTN",
    "tieu_luan": "Chấm tiểu luận hết học phần bậc DH",
    "huong_dan_khoa_luan": "H.Dẫn khóa luận tốt nghiệp bậc ĐH",
    "chu_tich_khoa_luan": "Chủ tịch HĐ chấm khóa luận bậc ĐH",
    "phan_bien_khoa_luan": "Đọc và phản biện khóa luận bậc ĐH",
    "thu_ky_khoa_luan": "Thư ký HĐ chấm khóa luận bậc ĐH",
    "huong_dan_ho_so_tn": "Hướng dẫn hồ sơ tốt nghiệp bậc ĐH",
    "cham_ho_so_tn": "Chấm hồ sơ tốt nghiệp bậc ĐH",
    "cham_tieu_luan_nckh": "Chấm tiểu luận nghiên cứu khoa học",
    "huong_dan_tieu_luan_nckh": "Hướng dẫn tiểu luận nghiên cứu khoa học",
    "cham_nckh_tich_luy": "Chấm sản phẩm nghiên cứu khoa học",
}

ACT_UNIT = {
    "ra_de": "đề+ĐA", "duyet_de": "đề",
    "tu_luan": "bài", "thuc_hanh": "bài", "tieu_luan": "bài",
    "huong_dan_khoa_luan": "khóa luận",
    "chu_tich_khoa_luan": "hội đồng",
    "phan_bien_khoa_luan": "hội đồng",
    "thu_ky_khoa_luan": "hội đồng",
    "huong_dan_ho_so_tn": "lớp",
    "cham_ho_so_tn": "hồ sơ",
    "cham_tieu_luan_nckh": "công trình",
    "huong_dan_tieu_luan_nckh": "công trình",
    "cham_nckh_tich_luy": "sản phẩm",
}

ACT_ORDER = [
    "ra_de", "duyet_de", "tu_luan", "thuc_hanh", "tieu_luan",
    "huong_dan_khoa_luan", "chu_tich_khoa_luan",
    "phan_bien_khoa_luan", "thu_ky_khoa_luan",
    "huong_dan_ho_so_tn",
    "cham_ho_so_tn",
    "cham_tieu_luan_nckh",
    "huong_dan_tieu_luan_nckh",
    "cham_nckh_tich_luy",
]

# ===== MAP CHỨC DANH CHỮ KÝ =====
FOOTER_LABEL = {
    "rade_chamthi": "TRƯỞNG PHÒNG KHẢO THÍ & ĐBCL",
    "khoa_luan": "TRƯỞNG PHÒNG KẾ HOẠCH TÀI CHÍNH",
}

# Column widths (dxa) — đo từ template gốc
COL_WIDTHS = [540, 604, 2293, 1143, 1141, 1334, 1150, 1334]
# Alignment cho từng cột data (8 cột)
COL_ALIGN  = ["center", "center", "left", "center", "right", "right", "right", "right"]
FONT_NAME  = "Times New Roman"


# =================== ĐỌC DỮ LIỆU ===================
def load_data_khoa_luan(xlsx_path: str) -> dict:
    wb = openpyxl.load_workbook(xlsx_path)
    ws = wb.active

    def _int(v): return int(v) if v else 0

    nam_hoc = ws["D1"].value or "2025 - 2026"
    hoc_ky  = str(ws["F1"].value or "2")

    teachers = []
    for row in ws.iter_rows(min_row=3, values_only=True):
        if row[0] is None:
            break
        teachers.append({
            "ma":                   str(row[0]).zfill(4),
            "ten":                  row[1],
            "co_thue":              str(row[2]).strip().upper() == "X" if row[2] else False,
            "huong_dan_ho_so_tn": _int(row[3]),
            "cham_ho_so_tn":   _int(row[4]),
            "cham_tieu_luan_nckh": _int(row[7]),
            "huong_dan_tieu_luan_nckh": _int(row[6]),
            "cham_nckh_tich_luy": _int(row[5])
        })

    teachers.sort(key=lambda x: x["ma"])
    return {"teachers": teachers, "nam_hoc": nam_hoc, "hoc_ky": hoc_ky}


def load_data_ho_so_tn_spnc(xlsx_path: str) -> dict:
    wb = openpyxl.load_workbook(xlsx_path)
    ws = wb.active

    def _int(v): return int(v) if v else 0

    nam_hoc = ws["D1"].value or "2025 - 2026"
    hoc_ky  = str(ws["F1"].value or "2")

    teachers = []
    for row in ws.iter_rows(min_row=3, values_only=True):
        if row[0] is None:
            break
        teachers.append({
            "ma":                   str(row[0]).zfill(4),
            "ten":                  row[1],
            "co_thue":              str(row[2]).strip().upper() == "X" if row[2] else False,
            "huong_dan_khoa_luan":  _int(row[3]),
            "chu_tich_khoa_luan":   _int(row[4]),
            "phan_bien_khoa_luan":  _int(row[5]),
            "thu_ky_khoa_luan":     _int(row[6])
        })

    teachers.sort(key=lambda x: x["ma"])
    return {"teachers": teachers, "nam_hoc": nam_hoc, "hoc_ky": hoc_ky}


def load_data(xlsx_path: str) -> dict:
    wb = openpyxl.load_workbook(xlsx_path)
    ws = wb.active

    def _int(v): return int(v) if v else 0

    nam_hoc = ws["D1"].value or "2025 - 2026"
    hoc_ky  = str(ws["F1"].value or "1")

    teachers = []
    for row in ws.iter_rows(min_row=3, values_only=True):
        if row[0] is None:
            break
        teachers.append({
            "ma":        str(row[0]).zfill(4),
            "ten":       row[1],
            "co_thue":   str(row[2]).strip().upper() == "X" if row[2] else False,
            "tu_luan":   _int(row[3]),
            "thuc_hanh": _int(row[4]),
            "tieu_luan": _int(row[5]),
            "ra_de":     _int(row[6]),
            "duyet_de":  _int(row[7]),
        })

    teachers.sort(key=lambda x: x["ma"])
    return {"teachers": teachers, "nam_hoc": nam_hoc, "hoc_ky": hoc_ky}


# =================== HELPERS ===================
def fmt_money(v: int) -> str:
    return f"{v:,.0f}đ".replace(",", ".")

def compute_acts(t: dict) -> list:
    rows = []
    for key in ACT_ORDER:
        qty = t.get(key, 0)
        if qty == 0:
            continue
        tien = qty * GIA[key]
        vat  = round(tien * 0.1) if t["co_thue"] else 0
        rows.append({
            "act":       ACT_NAME[key],
            "qty":       qty,
            "unit":      ACT_UNIT[key],
            "gia":       GIA[key],
            "tien":      tien,
            "vat":       vat,
            "thanh_tien": tien - vat,
        })
    return rows


# =================== XML BUILDERS ===================
def make_rPr(bold=False, sz=None):
    rPr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr in ["w:ascii", "w:hAnsi", "w:cs"]:
        fonts.set(qn(attr), FONT_NAME)
    rPr.append(fonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    if sz:
        sz_el = OxmlElement("w:sz")
        sz_el.set(qn("w:val"), str(sz))
        szCs_el = OxmlElement("w:szCs")
        szCs_el.set(qn("w:val"), str(sz))
        rPr.append(sz_el)
        rPr.append(szCs_el)
    return rPr

def make_pPr(align=None):
    pPr = OxmlElement("w:pPr")
    if align and align != "left":
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), align)
        pPr.append(jc)
    pPr.append(make_rPr())
    return pPr

def make_tc(text: str, align="left", bold=False,
            width_dxa=None, grid_span=None, sz=None):
    tc = OxmlElement("w:tc")

    tcPr = OxmlElement("w:tcPr")
    tcW  = OxmlElement("w:tcW")
    if width_dxa:
        tcW.set(qn("w:w"), str(width_dxa))
        tcW.set(qn("w:type"), "dxa")
    else:
        tcW.set(qn("w:w"), "0")
        tcW.set(qn("w:type"), "auto")
    tcPr.append(tcW)
    if grid_span and grid_span > 1:
        gs = OxmlElement("w:gridSpan")
        gs.set(qn("w:val"), str(grid_span))
        tcPr.append(gs)
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tcPr.append(va)
    tc.append(tcPr)

    p = OxmlElement("w:p")
    p.append(make_pPr(align))
    r = OxmlElement("w:r")
    r.append(make_rPr(bold, sz))
    t = OxmlElement("w:t")
    t.text = str(text)
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)
    tc.append(p)
    return tc

def make_tr(*cells):
    """cells: list of dict {text, align, bold, width, span, sz}"""
    tr = OxmlElement("w:tr")
    trPr = OxmlElement("w:trPr")
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:hRule"), "auto")
    trPr.append(trH)
    tr.append(trPr)
    for c in cells:
        tr.append(make_tc(
            text      = c.get("text", ""),
            align     = c.get("align", "left"),
            bold      = c.get("bold", False),
            width_dxa = c.get("width"),
            grid_span = c.get("span"),
            sz        = c.get("sz"),
        ))
    return tr


# =================== ROW FACTORIES ===================
def row_teacher(ma, ten):
    return make_tr({"text": f"♦ {ma} - {ten}", "span": 8})

def row_activity(stt, act, qty, unit, gia, tien, vat, thanh_tien):
    vals = [str(stt), "1", act, f"{qty} ({unit})",
            fmt_money(gia), fmt_money(tien), fmt_money(vat), fmt_money(thanh_tien)]
    return make_tr(*[{"text": v, "align": COL_ALIGN[i], "width": COL_WIDTHS[i]}
                     for i, v in enumerate(vals)])

def row_cong(subtotal, t_tax, net):
    """4 cells: span-5 'Cộng' | Tiền | Thuế | Thực lĩnh"""
    return make_tr(
        {"text": "Cộng",              "align": "center", "bold": True, "span": 5},
        {"text": fmt_money(subtotal), "align": "right",  "bold": True},
        {"text": fmt_money(t_tax),    "align": "right",  "bold": True},
        {"text": fmt_money(net),      "align": "right",  "bold": True},
    )

def row_grand(total, tax, net):
    return make_tr(
        {"text": "Cộng toàn bảng", "align": "center", "bold": True, "span": 5},
        {"text": fmt_money(total), "align": "right",  "bold": True},
        {"text": fmt_money(tax),   "align": "right",  "bold": True},
        {"text": fmt_money(net),   "align": "right",  "bold": True},
    )

def _replace_runs(p, new_text: str):
    """Xóa toàn bộ runs cũ, tạo lại 1 run duy nhất (giữ rPr gốc)."""
    first_r = p.find(qn("w:r"))
    old_rPr = None
    if first_r is not None:
        rp = first_r.find(qn("w:rPr"))
        if rp is not None: old_rPr = copy.deepcopy(rp)
    for r in p.findall(qn("w:r")): p.remove(r)
    new_r = OxmlElement("w:r")
    if old_rPr is None: old_rPr = OxmlElement("w:rPr")
    new_r.append(old_rPr)
    t_el = OxmlElement("w:t"); t_el.text = new_text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(t_el); p.append(new_r)

# =================== CẬP NHẬT TIÊU ĐỀ ===================
def update_header(doc, hoc_ky: str, nam_hoc: str):
    """Cập nhật dòng 'Hệ chính qui Học kỳ X, Năm học ...' trong header."""
    ht = doc.tables[0].rows[0].cells[0].tables[0]
    new_text = f"Hệ chính qui Học kỳ {hoc_ky}, Năm học {nam_hoc}"
    for row in ht.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if "Hệ chính qui" in para.text:
                    _replace_runs(para._p, new_text)
                    return


def update_footer(doc, report_type: str):
    """Cập nhật chức danh chữ ký (col0, row1) trong footer table."""
    label = FOOTER_LABEL.get(report_type, FOOTER_LABEL["rade_chamthi"])
    footer_table = doc.tables[0].rows[1].cells[0].tables[0]
    tc = footer_table.rows[1]._tr.findall(qn("w:tc"))[0]
    _replace_runs(tc.find(qn("w:p")), label)


# =================== GENERATE ===================
def generate_report(xlsx_path: str, template_path: str, report_type: str="rade_chamthi") -> str:
    if report_type == "rade_chamthi":
        data = load_data(xlsx_path)
    elif report_type == "khoa_luan":
        data = load_data_khoa_luan(xlsx_path)
    elif report_type == "hstn_spnc":
        data = load_data_ho_so_tn_spnc(xlsx_path)
    import json
    print(json.dumps(data, indent=3))
    teachers = data["teachers"]
    nam_hoc  = data["nam_hoc"]
    hoc_ky   = data["hoc_ky"]

    doc        = Document(template_path)
    outer_tbl  = doc.tables[0]
    data_table = outer_tbl.rows[0].cells[0].tables[1]
    tbl_elem   = data_table._tbl

    # Xóa toàn bộ dòng dữ liệu cũ, giữ lại header (row 0)
    for tr in tbl_elem.findall(qn("w:tr"))[1:]:
        tbl_elem.remove(tr)

    grand_total = grand_tax = 0

    for t in teachers:
        acts = compute_acts(t)
        if not acts:
            continue

        subtotal = sum(a["tien"] for a in acts)
        t_tax    = round(subtotal * 0.1) if t["co_thue"] else 0
        net      = subtotal - t_tax
        grand_total += subtotal
        grand_tax   += t_tax

        tbl_elem.append(row_teacher(t["ma"], t["ten"]))
        for stt, a in enumerate(acts, 1):
            tbl_elem.append(row_activity(
                stt, a["act"], a["qty"], a["unit"],
                a["gia"], a["tien"], a["vat"], a["thanh_tien"],
            ))
        tbl_elem.append(row_cong(subtotal, t_tax, net))

    tbl_elem.append(row_grand(grand_total, grand_tax, grand_total - grand_tax))

    # Cập nhật tiêu đề học kỳ / năm học
    update_header(doc, hoc_ky, nam_hoc)
    update_footer(doc, report_type)

    nam_hoc_clean = str(nam_hoc).replace(" ", "").replace("-", "_")
    output_path   = f"FIT_TTHDK_{report_type}_HK{hoc_ky}_{nam_hoc_clean}.docx"
    doc.save(output_path)

    # ===== SAVE TO MEMORY =====
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if __name__ == "__main__":
    generate_report(
        xlsx_path="Data.xlsx",
        template_path="Template.docx"
    )